import logging
import json
import os
import time
import vertexai
from pathlib import Path
from typing import Dict, Optional
from docx import Document
from io import BytesIO
from vertexai.generative_models import GenerativeModel
from sqlalchemy.orm import Session

# Configuración de logger
logger = logging.getLogger(__name__)

# Intentamos importar NER
try:
    from .ner_service import extract_entities
except ImportError:
    def extract_entities(text): return {}

class DOCXProcessor:
    def __init__(self):
        current_file = Path(__file__).resolve()
        root_dir = current_file.parent.parent.parent
        json_filename = "vertex_credenciales.json"
        path_to_keys = root_dir / json_filename

        self.project_id = "lector-pdfs-479603"
        self.location = "us-central1"

        if not path_to_keys.exists():
            logger.error(f"No se encontró el archivo '{json_filename}' en {root_dir}")
            self.model = None
            return

        try:
            os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = str(path_to_keys)
            vertexai.init(project=self.project_id, location=self.location)
            self.model = GenerativeModel("gemini-2.0-flash-001")
            logger.info(f"Vertex AI (DOCX) inicializado en el proyecto {self.project_id}")
        except Exception as e:
            logger.error(f"Error iniciando Vertex AI en DOCXProcessor: {e}")
            self.model = None

    def extract_text_from_docx(self, docx_bytes: bytes) -> str:
        """Extrae todo el texto plano del DOCX, incluyendo tablas."""
        try:
            doc = Document(BytesIO(docx_bytes))
            text = []
            
            # Extraer párrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extraer tablas (vital para el formato UPAO)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))
            
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error leyendo DOCX crudo: {e}")
            return ""

    def _parse_with_gemini(self, raw_text: str) -> Dict:
        """
        Usa Vertex AI para entender la estructura del Sílabo UPAO.
        """
        if not self.model:
            return {}

        prompt = f"""
        Actúa como un analista académico. Tienes el texto crudo de un Sílabo universitario.
        Extrae la información clave en un JSON estricto.

        Reglas:
        1. 'nombre': Nombre de la asignatura.
        2. 'codigo': Código del curso.
        3. 'ciclo': Número de ciclo (entero).
        4. 'descripcion': Texto de la SUMILLA.
        5. 'temas_clave': Lista de strings con temas técnicos (ej: ["Patrones", "SQL"]).
        6. 'texto_optimizado_sbert': Resumen para embeddings.

        Texto:
        {raw_text[:7000]}
        """

        max_retries = 3
        for attempt in range(max_retries):
            try:
                # Rate Limiting: Esperar un poco antes de cada intento
                time.sleep(2 + (attempt * 2)) 
                
                response = self.model.generate_content(prompt)
                cleaned_response = response.text.replace("```json", "").replace("```", "").strip()
                return json.loads(cleaned_response)
            except Exception as e:
                error_str = str(e)
                if "429" in error_str or "Resource exhausted" in error_str:
                    wait_time = (attempt + 1) * 10
                    logger.warning(f"⚠️ Quota excedida (429) en sílabo. Reintentando en {wait_time}s... ({attempt+1}/{max_retries})")
                    time.sleep(wait_time)
                else:
                    logger.error(f"Error parsing sílabo con Vertex AI: {e}")
                    return {}
        
        return {}

    def extract_syllabus_info(self, docx_bytes: bytes, filename: str = "") -> Dict:
        try:
            # 1. Leer texto crudo
            full_text = self.extract_text_from_docx(docx_bytes)
            if not full_text:
                return {'success': False, 'error': 'DOCX vacío o ilegible'}

            # 2. Interpretar con Vertex AI
            ai_data = self._parse_with_gemini(full_text)
            
            # Corrección para listas
            if isinstance(ai_data, list):
                ai_data = ai_data[0] if ai_data else {}

            # Valores por defecto
            nombre = ai_data.get('nombre')
            if not nombre and filename:
                nombre = filename.replace('.docx', '').replace('_', ' ')

            # 3. Enriquecer con NER local
            target_text = ai_data.get('texto_optimizado_sbert', full_text)
            entities = extract_entities(target_text)

            temas_gemini = ai_data.get('temas_clave', [])
            
            return {
                'success': True,
                'nombre': nombre,
                'codigo': ai_data.get('codigo'),
                'ciclo': ai_data.get('ciclo', 1),
                'descripcion': ai_data.get('descripcion', ''),
                'contenidos': list(set(temas_gemini + entities.get('contenidos', []))),
                'areas': entities.get('areas', []),
                'herramientas': entities.get('herramientas', []),
                'lenguajes': entities.get('lenguajes', []),
                'metodologias': entities.get('metodologias', []),
                'full_text': target_text,
                'raw_text_length': len(full_text)
            }
        except Exception as e:
            logger.error(f"Error procesando sílabo: {e}")
            return {'success': False, 'error': str(e)}

    def save_curso_to_db(self, db: Session, syllabus_info: Dict, drive_file_id: str) -> Optional[int]:
        try:
            from backend.database import crud
            existing = crud.get_curso_by_drive_id(db, drive_file_id)
            
            data = {
                "nombre": syllabus_info.get('nombre', 'Curso Desconocido'),
                "codigo": syllabus_info.get('codigo'),
                "ciclo": int(syllabus_info.get('ciclo', 1)),
                "descripcion": syllabus_info.get('descripcion'),
                "areas": syllabus_info.get('areas', []),
                "herramientas": syllabus_info.get('herramientas', []),
                "lenguajes": syllabus_info.get('lenguajes', []),
                "metodologias": syllabus_info.get('metodologias', []),
                "contenidos": syllabus_info.get('contenidos', []),
                "syllabus_text": syllabus_info.get('full_text', '')
            }

            if existing:
                crud.update_curso(db, existing.id, **data)
                return existing.id
            else:
                curso = crud.create_curso(db, drive_file_id=drive_file_id, **data)
                return curso.id
        except Exception as e:
            logger.error(f"Error BD Curso: {e}")
            return None

docx_processor = DOCXProcessor()